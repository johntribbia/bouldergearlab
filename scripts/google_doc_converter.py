#!/usr/bin/env python3
"""
Google Docs (via DOCX) to Markdown Converter
Converts .docx files to .md format with proper image extraction and linking.

Usage:
    converter = ReviewToMarkdownConverter("path/to/document.docx", output_dir="content/reviews/slug-name")
    converter.convert(
        title="Product Name Review",
        author="John Tribbia, Renee Krusemark",
        tags=["running", "shoes"],
        categories=["reviews"]
    )
"""

import os
import re
import datetime
import requests
import tempfile
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class ReviewToMarkdownConverter:
    def __init__(self, docx_path, output_dir="output"):
        """
        Initialize converter with docx file path or Google Docs URL.
        
        Args:
            docx_path: Path to the .docx file OR Google Docs URL
                      (e.g., https://docs.google.com/document/d/ABC123/edit)
            output_dir: Directory where markdown and images will be saved
        """
        self.output_dir = output_dir
        self.image_dir = output_dir
        
        # Check if it's a Google Docs URL
        if "docs.google.com" in str(docx_path):
            self.docx_path = self._download_google_doc(docx_path)
        else:
            self.docx_path = docx_path
            if not os.path.exists(self.docx_path):
                raise FileNotFoundError(f"Document not found: {self.docx_path}")
        
        self.doc = Document(self.docx_path)
        self.image_counter = 0
        self.image_map = {}  # Maps internal IDs to extracted filenames
        
        # Create output directories
        os.makedirs(self.output_dir, exist_ok=True)
    
    def _download_google_doc(self, url):
        """
        Download a Google Doc and convert it to DOCX format.
        
        Args:
            url: Google Docs URL (e.g., https://docs.google.com/document/d/ABC123/edit)
            
        Returns:
            str: Path to the downloaded DOCX file
        """
        print("📥 Downloading Google Doc...")
        
        # Extract document ID from URL
        match = re.search(r'/document/d/([a-zA-Z0-9-_]+)', url)
        if not match:
            raise ValueError("Invalid Google Docs URL. Use: https://docs.google.com/document/d/DOC_ID/edit")
        
        doc_id = match.group(1)
        export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=docx"
        
        try:
            response = requests.get(export_url, timeout=30)
            response.raise_for_status()
        except requests.exceptions.RequestException as e:
            raise RuntimeError(f"Failed to download document: {e}")
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        temp_file.write(response.content)
        temp_file.close()
        
        print(f"   ✓ Downloaded successfully")
        return temp_file.name
        
    def _extract_and_save_images(self):
        """
        Extracts all images from the .docx file in document order (top to bottom)
        and saves them as image_1.jpg, image_2.jpg, etc.
        
        Returns:
            dict: Mapping of relationship IDs to extracted filenames
        """
        image_map = {}
        seen_rids = set()
        
        # Iterate through paragraphs in document order to maintain image sequence
        for paragraph in self.doc.paragraphs:
            # Check for images in runs
            for run in paragraph.runs:
                for child in run._element:
                    if child.tag.endswith('drawing'):
                        # Found a drawing element, extract embedded images
                        for elem in child.iter():
                            if 'blip' in elem.tag:
                                embed_ref = elem.get(qn('r:embed'))
                                if embed_ref and embed_ref not in seen_rids:
                                    # This is a new image we haven't processed
                                    try:
                                        rel = self.doc.part.rels[embed_ref]
                                        if "image" in rel.target_ref:
                                            self.image_counter += 1
                                            new_name = f"image_{self.image_counter}.jpg"
                                            img_data = rel.target_part.blob
                                            
                                            img_path = os.path.join(self.image_dir, new_name)
                                            with open(img_path, "wb") as f:
                                                f.write(img_data)
                                            
                                            image_map[embed_ref] = new_name
                                            seen_rids.add(embed_ref)
                                            print(f"  ✓ Extracted: {new_name}")
                                    except (KeyError, AttributeError):
                                        pass
        
        self.image_map = image_map
        return image_map

    def _get_paragraph_style_name(self, paragraph):
        """Get the style name of a paragraph."""
        try:
            return paragraph.style.name
        except:
            return None

    def _is_heading(self, paragraph):
        """Detect if paragraph is a heading based on style or content."""
        style_name = self._get_paragraph_style_name(paragraph)
        text = paragraph.text.strip()
        
        # Check style-based headings
        if style_name and any(h in style_name for h in ["Heading 1", "Heading 2", "Heading 3"]):
            return True
        
        # Check formatting-based headings (bold, larger font, etc.)
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if first_run.bold and len(text) < 100:  # Short, bold text
                return True
        
        return False

    def _get_heading_level(self, paragraph):
        """Determine heading level (1-6) based on style."""
        style_name = self._get_paragraph_style_name(paragraph)
        
        if style_name:
            if "Heading 1" in style_name:
                return 1
            elif "Heading 2" in style_name:
                return 2
            elif "Heading 3" in style_name:
                return 3
        
        # Default to H3 for detected headings
        return 3

    def _extract_hyperlink_from_run(self, run):
        """
        Extract hyperlink URL and text from a run element.
        
        Args:
            run: A run object from python-docx
            
        Returns:
            tuple: (text, url) if hyperlink exists, (None, None) otherwise
        """
        try:
            # Look for hyperlink element in the run's parent (paragraph)
            # Hyperlinks are stored at paragraph level, not run level
            # We need to find the hyperlink that contains this run
            
            # First, check if this run is part of a hyperlink
            hyperlink = run._element.find('.//' + qn('w:hyperlink'))
            
            if hyperlink is not None:
                # Found a hyperlink element
                rId = hyperlink.get(qn('r:id'))
                
                if rId:
                    # Get the target URL from relationships
                    try:
                        rel = self.doc.part.rels[rId]
                        url = rel.target_ref
                        text = run.text
                        return text, url
                    except (KeyError, AttributeError):
                        pass
            
            # Check if run itself is inside a hyperlink in the paragraph
            parent_hyperlink = None
            current = run._element
            
            # Walk up the XML tree to find parent hyperlink
            while current is not None:
                if current.tag == qn('w:hyperlink'):
                    parent_hyperlink = current
                    break
                current = current.getparent()
            
            if parent_hyperlink is not None:
                rId = parent_hyperlink.get(qn('r:id'))
                if rId:
                    try:
                        rel = self.doc.part.rels[rId]
                        url = rel.target_ref
                        text = run.text
                        return text, url
                    except (KeyError, AttributeError):
                        pass
        except:
            pass
        
        return None, None

    def _process_paragraph_text(self, paragraph):
        """
        Process paragraph text and extract inline images and hyperlinks.
        
        Returns:
            list: Lines of processed text (may include image references and hyperlinks)
        """
        lines = []
        para_text = ""
        
        for run in paragraph.runs:
            # Check if run contains an image (drawing object)
            for child in run._element:
                if child.tag.endswith('drawing'):
                    # Found an inline image
                    if para_text.strip():
                        lines.append(para_text.strip())
                        para_text = ""
                    
                    # Extract image from drawing
                    for inline_shape in child.iter():
                        for pic in inline_shape.iter():
                            if 'blip' in pic.tag:
                                # Image found - get embed reference
                                embed_ref = pic.get(qn('r:embed'))
                                if embed_ref and embed_ref in self.image_map:
                                    img_filename = self.image_map[embed_ref]
                                    lines.append(f"\n![{img_filename}]({img_filename})\n")
                    continue
            
            # Check for hyperlinks in this run
            link_text, link_url = self._extract_hyperlink_from_run(run)
            
            if link_text and link_url:
                # Add any preceding text first
                if para_text.strip():
                    lines.append(para_text.strip())
                    para_text = ""
                # Add hyperlink as markdown
                para_text += f"[{link_text}]({link_url})"
            else:
                # Regular text
                para_text += run.text
        
        # Add remaining text
        if para_text.strip():
            lines.append(para_text.strip())
        
        return lines

    def _format_text_with_markdown(self, text):
        """
        Convert basic formatting to Markdown.
        
        Args:
            text: Text to format
            
        Returns:
            str: Text with Markdown formatting applied
        """
        # This is a simplified version - DOCX formatting is complex
        # The python-docx library handles some of this, but inline formatting
        # is better handled manually if needed
        return text

    def convert(self, title, author, tags=None, categories=None):
        """
        Convert DOCX to Markdown format suitable for Boulder Gear Lab.
        
        Args:
            title: Review title (e.g., "Hoka Speedgoat 5 GTX Review")
            author: Author name(s) (e.g., "John Tribbia, Renee Krusemark")
            tags: List of tags (e.g., ["running", "shoes"])
            categories: List of categories (default: ["reviews"])
        """
        if tags is None:
            tags = []
        if categories is None:
            categories = ["reviews"]
        
        print(f"\n📝 Converting: {title}")
        print("=" * 60)
        
        # Step 1: Extract Images
        print("\n1. Extracting images...")
        self._extract_and_save_images()
        
        # Step 2: Build Front Matter
        print("2. Building front matter...")
        today = datetime.date.today().strftime("%Y-%m-%d")
        
        markdown_lines = [
            "---",
            f'title: "{title}"',
            f'date: {today}',
            'banner: "image_1.jpg"',
            f'tags: {tags}',
            f'categories: {categories}',
            'description: ""',
            'draft: false',
            "---",
            "<!--more-->",
            "",
            f"*Article by {author}*",
            "",
            "Original Post from RoadTrailRun",
            "([link](https://www.roadtrailrun.com/))",
            "",
            '<a href="https://www.roadtrailrun.com"',
            'class="button primary button-wrapper"><span>Read All RoadTrailRun',
            'Reviews Here</span></a>',
            "",
        ]
        
        # Step 3: Process Document Content with Inline Images
        print("3. Processing document content...")
        
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            
            # Skip empty paragraphs
            if not text and not any(run._element.find('.//' + qn('w:drawing')) for run in paragraph.runs):
                continue
            
            # Check for headings
            if self._is_heading(paragraph):
                level = self._get_heading_level(paragraph)
                markdown_lines.append(f"\n{'#' * level} {text}\n")
                continue
            
            # Process runs in order to handle inline images, text, and hyperlinks
            para_content = []
            for run in paragraph.runs:
                # Check for images in this run
                has_image = False
                for child in run._element:
                    if child.tag.endswith('drawing'):
                        has_image = True
                        # Extract and link the image
                        for elem in child.iter():
                            if 'blip' in elem.tag:
                                embed_ref = elem.get(qn('r:embed'))
                                if embed_ref and embed_ref in self.image_map:
                                    img_filename = self.image_map[embed_ref]
                                    # Add any preceding text first
                                    if para_content:
                                        markdown_lines.append("".join(para_content).strip())
                                        para_content = []
                                    # Add the image
                                    markdown_lines.append(f"\n![{img_filename}]({img_filename})\n")
                
                # Handle text content - check for hyperlinks
                if not has_image and run.text:
                    # Check if this run contains a hyperlink
                    link_text, link_url = self._extract_hyperlink_from_run(run)
                    
                    if link_text and link_url:
                        # Hyperlink found - format as markdown link
                        para_content.append(f"[{link_text}]({link_url})")
                    else:
                        # Regular text
                        para_content.append(run.text)
            
            # Add any remaining text content
            if para_content:
                text_content = "".join(para_content).strip()
                if text_content:
                    markdown_lines.append(text_content)
        
        # Step 4: Save Markdown File
        print("4. Saving markdown file...")
        
        # Generate slug from title
        slug = title.lower()
        slug = re.sub(r'[^\w\s-]', '', slug)  # Remove special chars
        slug = re.sub(r'[-\s]+', '-', slug)  # Replace spaces/hyphens with single hyphen
        slug = slug.strip('-')
        
        file_name = "index.md"
        file_path = os.path.join(self.output_dir, file_name)
        
        with open(file_path, "w", encoding="utf-8") as f:
            f.write("\n".join(markdown_lines))
        
        print(f"   ✓ Created: {file_path}")
        print("\n" + "=" * 60)
        print(f"✅ Conversion complete!")
        print(f"   Files saved to: {self.output_dir}")
        print(f"   Total images extracted: {self.image_counter}")
        print(f"   Markdown file: {file_name}")
        
        return file_path


def main():
    """Example usage - uncomment and modify as needed."""
    
    # Example 1: Convert a review
    # converter = ReviewToMarkdownConverter(
    #     docx_path="~/Downloads/My_Product_Review.docx",
    #     output_dir="content/reviews/my-product-review"
    # )
    # 
    # converter.convert(
    #     title="Product Name Review",
    #     author="John Tribbia, Renee Krusemark",
    #     tags=["running", "shoes"],
    #     categories=["reviews"]
    # )


if __name__ == "__main__":
    main()
